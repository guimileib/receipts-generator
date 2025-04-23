from io import BytesIO
import os
from functools import wraps

from flask import Blueprint, request, render_template, redirect, session, flash, send_file, url_for
from dotenv import load_dotenv
import pandas as pd
from docx import Document

from . import db
from .clientes import Cliente
from .pdf import gerar_pdf
from .mail import enviar_email

main = Blueprint('main', __name__, template_folder=os.path.join(os.path.dirname(__file__), '..', 'templates'))

def login_required(f):
    @wraps(f) 
    def decorated_function(*args, **kwargs):
        # Verificação mais robusta da sessão
        if not session.get('logged_in', False):
            # Armazenar URL solicitada para redirecionamento após login
            session['next_url'] = request.url
            flash('Você precisa fazer login primeiro', 'warning')
            return redirect(url_for('main.login'))
        return f(*args, **kwargs)
    return decorated_function

@main.route('/')
def index():
    return redirect(url_for('main.login'))

@main.route("/login", methods=['GET', 'POST'])
def login():
    # Se já estiver logado, redirecione para o dashboard
    if session.get('logged_in', False):
        return redirect(url_for('main.dashboard'))
        
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        load_dotenv()

        if username == os.getenv('USER') and password == os.getenv('PASSWORD'):
            # Definir a sessão de forma explícita
            session['logged_in'] = True
            session['username'] = username
            
            # Verificar se há uma URL para redirecionamento
            next_url = session.pop('next_url', None)
            
            flash('Login efetuado com sucesso!', 'success')
            if next_url:
                return redirect(next_url)
            return redirect(url_for('main.dashboard'))
        else:
            flash('Usuário ou senha inválidos', 'danger')
    
    return render_template('login.html')

@main.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html')

@main.route('/logout')
def logout():
    # Remover itens específicos da sessão ao invés de limpar tudo
    session.pop('logged_in', None)
    session.pop('username', None)
    flash('Você saiu da conta!', 'info')
    return redirect(url_for('main.login'))

@main.route('/contrato', methods=['GET'])
@login_required
def contrato():
    return render_template('contrato.html')

def substituir_texto_em_documento(doc, substitutos):
    """
    Substitui texto em um documento Word, lidando com texto dividido entre múltiplos runs
    """
    # Para cada parágrafo no documento
    for paragrafo in doc.paragraphs:
        # Para cada chave e valor no dicionário de substituições
        for chave, valor in substitutos.items():
            # Se a chave estiver no texto do parágrafo
            if chave in paragrafo.text:
                texto_inline = paragrafo.text
                texto_inline = texto_inline.replace(chave, valor)
                
                # Limpar todos os runs existentes e adicionar um novo com o texto substituído
                for i in range(len(paragrafo.runs)):
                    paragrafo.runs[0]._element.getparent().remove(paragrafo.runs[0]._element)
                
                run = paragrafo.add_run(texto_inline)
                # Nota: isso remove a formatação original

    # Também verificar tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for chave, valor in substitutos.items():
                        if chave in paragrafo.text:
                            texto_inline = paragrafo.text
                            texto_inline = texto_inline.replace(chave, valor)
                            
                            for i in range(len(paragrafo.runs)):
                                paragrafo.runs[0]._element.getparent().remove(paragrafo.runs[0]._element)
                            
                            run = paragrafo.add_run(texto_inline)
    
@main.route('/contrato', methods=['POST'])
@login_required
def gerar_contrato():
    if not session.get('logged_in'):
        flash('Você precisa fazer login para gerar contratos', 'warning')
        return redirect(url_for('main.login'))
    
    nome = request.form.get('nome', '')
    cpf = request.form.get('cpf', '')
    servico = request.form.get('servico', '')
    valor = request.form.get('valor', '')
    endereco = request.form.get('endereco', '')
    data = request.form.get('data', '')
    
    caminho_modelo = os.path.join(os.getcwd(), "static", "doc", "contrato_modelo.docx")
    
    # Verificar se o arquivo existe
    if not os.path.exists(caminho_modelo):
        flash(f'Arquivo de modelo não encontrado em: {caminho_modelo}', 'danger')
        return redirect(url_for('main.contrato'))
    
    try:
        doc = Document(caminho_modelo)
        
        substitutos = {
            "{nome}": nome,
            "{cpf}": cpf,
            "{servico}": servico,
            "{valor}": valor,
            "{endereco}": endereco,
            "{data}": data
        }

        substituir_texto_em_documento(doc, substitutos)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        nome_arquivo = f"contrato_{nome.replace(' ', '_')}.docx"
        
        return send_file(
            output,
            as_attachment=True,
            download_name=nome_arquivo,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        flash(f'Erro ao gerar contrato: {str(e)}', 'danger')
        return redirect(url_for('main.contrato'))

@main.route('/register', methods=['GET'])
@login_required
def show_register_form():
    return render_template('register.html')

@main.route('/register', methods=['POST'])
@login_required
def register():
    try:
        nome = request.form['nome']
        email = request.form['email']
        telefone = request.form['telefone']
        endereco = request.form['endereco']
        valor_servico = float(request.form['service_value'])
        descricao_servico = request.form['service_description']

        novo_cliente = Cliente(
            nome=nome, 
            email=email, 
            telefone=telefone, 
            valor_servico=valor_servico, 
            descricao_servico=descricao_servico, 
            endereco=endereco
        )

        db.session.add(novo_cliente)
        db.session.commit()

        # Gerar PDF e enviar email
        pdf_buffer = gerar_pdf(novo_cliente)
        enviar_email(novo_cliente, pdf_buffer)
        
        flash('Cliente cadastrado e e-mail enviado!', 'success')
        return redirect(url_for('main.listar_clientes'))
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao cadastrar cliente: {str(e)}', 'danger')
        return redirect(url_for('main.show_register_form'))

@main.route('/clientes')
@login_required
def listar_clientes():
    clientes = Cliente.query.all()
    return render_template('clientes.html', clientes=clientes)

@main.route('/exportar')
@login_required
def exportar():
    try:
        clientes = Cliente.query.all()
        clientes_lista = [{
            'Nome': cliente.nome,
            'Email': cliente.email,
            'Telefone': cliente.telefone,
            'Descrição do Serviço': cliente.descricao_servico,
            'Valor do Serviço': cliente.valor_servico,
            'Endereço': cliente.endereco
        } for cliente in clientes]

        df = pd.DataFrame(clientes_lista)
        output = BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        
        output.seek(0)
        
        return send_file(
            output, 
            download_name='clientes.xlsx', 
            as_attachment=True, 
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    except Exception as e:
        flash(f'Erro ao exportar dados: {str(e)}', 'danger')
        return redirect(url_for('main.listar_clientes'))
    
@main.route('/excluir_cliente/<int:id>', methods=['POST'])
@login_required
def excluir_cliente(id):
    try:
        cliente = Cliente.query.get_or_404(id)
        db.session.delete(cliente)
        db.session.commit()
        flash('Cliente excluído com sucesso!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao excluir cliente: {str(e)}', 'danger')
    
    return redirect(url_for('main.listar_clientes'))
    